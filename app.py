import streamlit as st
import io
import os
import fitz  # PyMuPDF: Thư viện xử lý PDF siêu nhanh, không cần Poppler!
from PIL import Image
from docx import Document
from google import genai
from google.genai import types

# ---------------------------------------------------------
# 1. CẤU HÌNH TRANG WEB STREAMLIT
# ---------------------------------------------------------
st.set_page_config(
    page_title="Chuyển đổi Tài liệu & PDF Scan AI",
    page_icon="📄",
    layout="wide"
)

st.title("📄 OCR & Tái Tạo Form (Đã hỗ trợ PDF Scan nhiều trang)")
st.caption("Tự động chuyển đổi File Ảnh & PDF Scan thành Word/Excel, giữ nguyên cấu trúc bảng biểu và form mẫu.")

# ---------------------------------------------------------
# 2. CẤU HÌNH GEMINI API KEY MẶC ĐỊNH / TÍCH HỢP SẴN
# ---------------------------------------------------------
# Điền API Key mặc định của bạn vào đây nếu muốn hardcode sẵn (hoặc dùng st.secrets)
DEFAULT_API_KEY = st.secrets.get("GEMINI_API_KEY", os.environ.get("GEMINI_API_KEY", ""))

with st.sidebar:
    st.header("⚙️ Cấu hình Hệ thống")
    user_api_key = st.text_input(
        "Nhập API Key riêng (Tùy chọn):", 
        type="password",
        help="Hệ thống đã tích hợp sẵn API Key miễn phí. Bạn chỉ cần nhập nếu muốn dùng Key riêng của mình."
    )
    
    # Ưu tiên lấy key do người dùng nhập, nếu không có thì dùng key mặc định tích hợp sẵn
    active_api_key = user_api_key.strip() if user_api_key.strip() else DEFAULT_API_KEY
    
    if user_api_key.strip():
        st.success("🔑 Đang dùng API Key riêng của bạn")
    elif DEFAULT_API_KEY:
        st.info("🎁 Đang dùng API Key tích hợp sẵn (Miễn phí)")
    else:
        st.warning("⚠️ Chưa cấu hình API Key mặc định. Vui lòng nhập API Key để sử dụng.")
        
    st.markdown("[Lấy API Key miễn phí tại đây](https://aistudio.google.com/)")
    st.divider()
    st.info("💡 **Hỗ trợ PDF Scan:** Hệ thống tự động tách từng trang PDF thành ảnh độ phân giải cao (300 DPI) để AI nhận diện chuẩn xác nhất.")

# ---------------------------------------------------------
# 3. CÁC HÀM XỬ LÝ LÕI (CORE FUNCTIONS)
# ---------------------------------------------------------

def pdf_to_images(pdf_bytes):
    """
    Chuyển đổi toàn bộ các trang PDF thành danh sách các hình ảnh (Bytes) 
    chất lượng cao dùng PyMuPDF (Không phụ thuộc Poppler).
    """
    images = []
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    
    for page_num in range(len(doc)):
        page = doc[page_num]
        # Zoom 2.0x ~ 300 DPI giúp AI đọc chữ nhỏ/mờ trên PDF Scan tốt hơn hẳn
        pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0))
        img_bytes = pix.tobytes("jpeg")
        images.append((page_num + 1, img_bytes))
        
    doc.close()
    return images

def process_image_ocr(image_bytes, api_key):
    """
    Gửi ảnh sang Gemini Vision để phân tích và trích xuất bố cục form/bảng.
    """
    client = genai.Client(api_key=api_key)
    
    prompt = """
    Hãy đóng vai là một chuyên gia OCR và tái tạo tài liệu hàng đầu:
    1. Trích xuất toàn bộ văn bản từ hình ảnh này (đặc biệt hỗ trợ tiếng Việt chuẩn xác).
    2. Nhận diện cấu trúc của tài liệu (tiêu đề, các đoạn văn, bảng biểu, các trường thông tin điền form).
    3. Trả về kết quả dưới dạng Markdown chuẩn:
       - Dùng dấu # cho Tiêu đề.
       - Dùng định dạng Bảng Markdown (| Cột 1 | Cột 2 |) cho bất kỳ bảng biểu hoặc form điền thông tin nào.
       - Giữ đúng thứ tự và bố cục từ trên xuống dưới của tài liệu gốc.
    Không thêm bất kỳ lời giới thiệu nào, chỉ trả về nội dung tài liệu dưới dạng Markdown.
    """
    
    # Đổi tên model thành 'gemini-2.0-flash' (hoặc 'gemini-1.5-flash-latest')
    response = client.models.generate_content(
        model='gemini-1.5-flash',
        contents=[
            types.Part.from_bytes(data=image_bytes, mime_type="image/jpeg"),
            prompt
        ]
    )
    return response.text
def export_to_word(markdown_text):
    """
    Chuyển đổi Markdown thành file MS Word (.docx) bảo toàn Bảng (Table Grid).
    """
    doc = Document()
    lines = markdown_text.split('\n')
    
    table_data = []
    in_table = False

    for line in lines:
        stripped = line.strip()
        
        # Xử lý Bảng biểu Markdown
        if stripped.startswith('|') and stripped.endswith('|'):
            in_table = True
            if '---' in stripped:
                continue
            row_data = [cell.strip() for cell in stripped.split('|')[1:-1]]
            table_data.append(row_data)
        else:
            if in_table and table_data:
                cols_count = max(len(r) for r in table_data)
                table = doc.add_table(rows=len(table_data), cols=cols_count)
                table.style = 'Table Grid'
                for r_idx, row in enumerate(table_data):
                    for c_idx, val in enumerate(row):
                        if c_idx < cols_count:
                            table.cell(r_idx, c_idx).text = val
                table_data = []
                in_table = False
            
            if stripped.startswith('# '):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith('## '):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith('### '):
                doc.add_heading(stripped[4:], level=3)
            elif stripped:
                doc.add_paragraph(stripped)

    if in_table and table_data:
        cols_count = max(len(r) for r in table_data)
        table = doc.add_table(rows=len(table_data), cols=cols_count)
        table.style = 'Table Grid'
        for r_idx, row in enumerate(table_data):
            for c_idx, val in enumerate(row):
                if c_idx < cols_count:
                    table.cell(r_idx, c_idx).text = val

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ---------------------------------------------------------
# 4. GIAO DIỆN NGUỜI DÙNG (USER INTERFACE)
# ---------------------------------------------------------

uploaded_files = st.file_uploader(
    "Tải lên các file ảnh chụp (JPG, PNG) hoặc File PDF Scan:", 
    type=["jpg", "jpeg", "png", "pdf"], 
    accept_multiple_files=True
)

if uploaded_files:
    st.success(f"Đã chọn {len(uploaded_files)} file.")
    
    if 'ocr_results' not in st.session_state:
        st.session_state['ocr_results'] = {}

    if st.button("🚀 Bắt đầu Nhận diện & Tái tạo (OCR)", type="primary"):
        if not active_api_key:
            st.error("Chưa tìm thấy Gemini API Key! Vui lòng cấu hình key mặc định hoặc nhập key ở thanh bên trái.")
        else:
            status_progress = st.empty()
            
            for file_idx, file in enumerate(uploaded_files):
                file_bytes = file.read()
                
                # --- XỬ LÝ FILE ẢNH (JPG / PNG) ---
                if file.type.startswith("image/"):
                    status_progress.info(f"⏳ Đang xử lý file ảnh: **{file.name}**...")
                    try:
                        res_text = process_image_ocr(file_bytes, active_api_key)
                        st.session_state['ocr_results'][file.name] = res_text
                    except Exception as e:
                        st.error(f"Lỗi khi OCR file {file.name}: {e}")
                
                # --- XỬ LÝ FILE PDF SCAN (NHIỀU TRANG) ---
                elif file.type == "application/pdf":
                    status_progress.info(f"⏳ Đang trích xuất các trang từ PDF Scan: **{file.name}**...")
                    try:
                        pages = pdf_to_images(file_bytes)
                        full_pdf_text = []
                        
                        for page_num, img_bytes in pages:
                            status_progress.info(f"⏳ Đang OCR file **{file.name}** (Trang {page_num}/{len(pages)})...")
                            page_text = process_image_ocr(img_bytes, active_api_key)
                            
                            # Lưu từng trang riêng biệt
                            st.session_state['ocr_results'][f"{file.name} - Trang {page_num}"] = page_text
                            full_pdf_text.append(f"<!-- TRANG {page_num} -->\n" + page_text)
                        
                        # Nếu PDF có nhiều trang, tự tạo thêm 1 bản gộp toàn bộ các trang
                        if len(pages) > 1:
                            st.session_state['ocr_results'][f"{file.name} - [TOÀN BỘ FILE]"] = "\n\n---\n\n".join(full_pdf_text)
                            
                    except Exception as e:
                        st.error(f"Lỗi khi xử lý PDF {file.name}: {e}")
            
            status_progress.success("✅ Hoàn tất nhận diện toàn bộ tài liệu!")

    # ---------------------------------------------------------
    # 5. XEM TRƯỚC, CHỈNH SỬA VÀ XUẤT TÀI LIỆU
    # ---------------------------------------------------------
    if st.session_state['ocr_results']:
        st.divider()
        st.subheader("📝 Xem trước, Chỉnh sửa & Xuất Tài liệu")
        
        selected_file_name = st.selectbox(
            "Chọn trang/file cần xem trước và xuất kết quả:", 
            options=list(st.session_state['ocr_results'].keys())
        )
        
        col1, col2 = st.columns([1, 1])
        
        with col1:
            st.markdown("**Trình soạn thảo văn bản (Sửa trực tiếp trước khi xuất):**")
            edited_text = st.text_area(
                label="Bản chỉnh sửa", 
                value=st.session_state['ocr_results'][selected_file_name], 
                height=450,
                key=f"editor_{selected_file_name}"
            )
            st.session_state['ocr_results'][selected_file_name] = edited_text

        with col2:
            st.markdown("**Kết quả hiển thị xem trước (Rendered Layout):**")
            st.markdown(edited_text)

        st.divider()
        st.subheader("📥 Xuất file tài liệu")
        
        col_docx, col_txt = st.columns(2)
        
        with col_docx:
            word_file = export_to_word(edited_text)
            st.download_button(
                label="📄 Tải về file Word (.docx)",
                data=word_file,
                file_name=f"{selected_file_name.replace(' ', '_')}_OCR.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
        with col_txt:
            st.download_button(
                label="📝 Tải về file Text (.txt)",
                data=edited_text,
                file_name=f"{selected_file_name.replace(' ', '_')}_OCR.txt",
                mime="text/plain"
            )
